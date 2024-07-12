import docx
import pandas as pd
import pathlib
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
import torch
from transformers import pipeline

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def load_all_files(directory_path):
    data = []
    for file_path in pathlib.Path(directory_path).glob("*"):
        if file_path.suffix == '.csv':
            df = pd.read_csv(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
        elif file_path.suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                data.append(Document(page_content=content))
        elif file_path.suffix == '.docx':
            content = read_docx(file_path)
            data.append(Document(page_content=content))
        elif file_path.suffix == '.xlsx':
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
    return data

def interpret_files(documents):
    print(f"Total documents loaded: {len(documents)}")
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    texts = splitter.split_documents(documents)
    return texts

def create_embeddings():
    print("Creating embeddings")
    model_name = "sentence-transformers/all-MiniLM-L6-v2"  # Example model
    embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs={'device': 'cuda' if torch.cuda.is_available() else 'cpu'}
    )
    return embeddings

def save(texts, embeddings):
    print("Saving data")
    # Store data into database
    db = Chroma.from_documents(texts, embedding=embeddings, persist_directory="test_index")
    db.persist()

def retrieve_answer(question, retriever, qa_pipeline):
    search_results = retriever.similarity_search_with_score(question, k=5)
    answers = []
    for doc, _ in search_results:
        answer = qa_pipeline(question=question, context=doc.page_content)
        answers.append((answer['answer'], answer['score']))
    return sorted(answers, key=lambda x: x[1], reverse=True)[0][0]

# Example usage
directory_path = 'D:\Axis-FAQ-chatbot\Data'  # Replace with your directory path
documents = load_all_files(directory_path)
texts = interpret_files(documents)
embeddings = create_embeddings()
save(texts, embeddings)

# Set up retriever
db = Chroma(persist_directory="test_index", embedding_function=embeddings)
retriever = db.as_retriever()

# Set up RoBERTa question-answering pipeline
qa_pipeline = pipeline("question-answering", model="deepset/roberta-base-squad2")

# Example question
question = "What is the interest rate for savings accounts?"
answer = retrieve_answer(question, retriever, qa_pipeline)
print(f"Answer: {answer}")
