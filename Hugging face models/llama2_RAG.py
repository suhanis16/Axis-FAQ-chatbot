# from langchain.document_loaders import DirectoryLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import CTransformers
from langchain import PromptTemplate
from langchain.chains import RetrievalQA

from IPython.display import display, HTML
import json
import time
import pathlib

# from langchain.document_loaders import CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import pandas as pd
from langchain.docstore.document import Document


import pandas as pd
import pathlib
import docx
from langchain.docstore.document import Document
import os

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def load_all_files(directory_path):
    data = []
    for file_path in pathlib.Path(directory_path).glob("*"):
        if file_path.suffix == '.csv':
            df = pd.read_csv(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
        elif file_path.suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                data.append(Document(page_content=content))
        elif file_path.suffix == '.docx':
            content = read_docx(file_path)
            data.append(Document(page_content=content))
        elif file_path.suffix == '.xlsx':
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
    return data


def interpret_files(documents):
    # Interpret information in the CSV file
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    texts = splitter.split_documents(documents)
    return texts

def create_embeddings():
    # Create embeddings
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    return embeddings

def save(texts,embeddings):
    # Create and save the local database
    db = FAISS.from_documents(texts, embeddings)
    db.save_local("faiss")

def load_llm(model_name,model_path):
    if model_name == "llama":
        # load the language model
        llm = CTransformers(model=model_path,
                            model_type='llama',
                            config={'max_new_tokens': 256, 'temperature': 0.01})
        
    return llm
                            

# load the interpreted information from the local database
# embeddings = HuggingFaceEmbeddings(
#     model_name="sentence-transformers/all-MiniLM-L6-v2",
#     model_kwargs={'device': 'cpu'})

def retrieve_docs(embeddings,llm):
    # prepare the template we will use when prompting the AI
    template = """Use the following pieces of information to answer the user's question.
    If you don't know the answer, just say that you don't know, don't try to make up an answer.
    Context: {context}
    Question: {question}
    Only return the helpful answer below and nothing else.
    Helpful answer:
    """

    db = FAISS.load_local("faiss", embeddings,allow_dangerous_deserialization=True)

    # prepare a version of the llm pre-loaded with the local content
    retriever = db.as_retriever(search_kwargs={'k': 2})
    prompt = PromptTemplate(
        template=template,
        input_variables=['context', 'question'])

    QA_LLM = RetrievalQA.from_chain_type(llm=llm,
                                        chain_type='stuff',
                                        retriever=retriever,
                                        return_source_documents=True,
                                        chain_type_kwargs={'prompt': prompt})
    
    return QA_LLM


def query(model, question):
    model_path = model.combine_documents_chain.llm_chain.llm.model
    model_name = pathlib.Path(model_path).name
    time_start = time.time()
    output = model({'query': question})
    response = output["result"]
    time_elapsed = time.time() - time_start
    print(f'<code>{model_name} response time: {time_elapsed:.02f} sec</code>')
    print(f'<strong>Question:</strong> {question}')
    print(f'<strong>Answer:</strong> {response}')


def main():

    data_path = "D:\Axis-FAQ-chatbot\Data"
    documents = load_all_files(data_path)
    texts = interpret_files(documents)
    embeddings = create_embeddings()
    save(texts,embeddings)

    model_path = "D:\Axis\models\llama-2-7b-chat.ggmlv3.q8_0.bin"

    llm = load_llm("llama", model_path)
    QA_LLM = retrieve_docs(embeddings,llm)

    user_input = input("What is your question? \n")

    print("Retrieving answer...")
    # Show running time
    # print(time.time())

    # print(query(QA_LLM,user_input))


    return query(QA_LLM,user_input)


if __name__=="__main__":
    main()


