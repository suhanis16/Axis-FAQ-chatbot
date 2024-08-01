import warnings
warnings.filterwarnings("ignore")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from langchain import HuggingFacePipeline
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
import torch
import pandas as pd
import pathlib
import docx
from langchain.docstore.document import Document
import time

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def load_all_files(directory_path):
    data = []
    for file_path in pathlib.Path(directory_path).glob("*"):
        if file_path.suffix == '.csv':
            df = pd.read_csv(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
        elif file_path.suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                data.append(Document(page_content=content))
        elif file_path.suffix == '.docx':
            content = read_docx(file_path)
            data.append(Document(page_content=content))
        elif file_path.suffix == '.xlsx':
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
    return data


def interpret_files(documents):
    print(f"Total documents loaded: {len(documents)}")
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    texts = splitter.split_documents(documents)
    # print(f"Total texts generated: {len(texts)}")
    return texts

def create_embeddings():
    print("Creating embeddings")
    if torch.cuda.is_available():
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cuda'}
        )
    else:
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
    return embeddings

def save(texts, embeddings):
    print("Saving data")
    # Store data into database
    db=Chroma.from_documents(texts,embedding=embeddings,persist_directory="test_index")
    db.persist()
    # db = FAISS.from_documents(texts, embeddings)
    # db.save_local("faiss")


def load_llm(model_name):
    print("Loading LLM")
    if model_name == "phi3":
        print("Loading Phi3 model")
        tokenizer = AutoTokenizer.from_pretrained("microsoft/Phi-3-mini-4k-instruct")
        model = AutoModelForCausalLM.from_pretrained("microsoft/Phi-3-mini-4k-instruct", device_map='auto', torch_dtype="auto", trust_remote_code=True,)

    elif model_name == "llama":
        print("Loading Llama model")
        tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-2-7b-chat-hf")
        model = AutoModelForCausalLM.from_pretrained("meta-llama/Llama-2-7b-chat-hf", device_map='auto', torch_dtype=torch.float16, trust_remote_code=True,)
        
    elif model_name == "gemma":
        print("Loading gemma model")
        tokenizer = AutoTokenizer.from_pretrained("google/gemma-2-9b-it")
        model = AutoModelForCausalLM.from_pretrained("google/gemma-2-9b-it", device_map='auto', torch_dtype=torch.float16, trust_remote_code=True,)

    pipe = pipeline("text-generation", model=model, tokenizer=tokenizer, max_new_tokens=300)
    llm = HuggingFacePipeline(pipeline=pipe)
    return llm


def retrieve_docs(embeddings, llm):
    print("Retrieving documents")

    # Load the database
    vectordb = Chroma(persist_directory="test_index", embedding_function = embeddings)

    # Load the retriver
    retriever = vectordb.as_retriever(search_kwargs = {"k" : 2})

    print("Retrieved documents")

    qna_prompt_template= """Use the following pieces of information to answer the user's question.
    If you don't know the answer, just say that you don't know, don't try to make up an answer.
    {context}
    Question: {question}
    Answer:"""

    PROMPT = PromptTemplate(
       template=qna_prompt_template, input_variables=["context","question"] # Change 'Context' to 'context' and 'Question' to 'question'
    )

    print("Sending the chain")
    # Define the QNA chain
    chain = RetrievalQA.from_chain_type(llm=llm,
                                             chain_type='stuff',
                                             retriever=retriever,
                                             return_source_documents=True,
                                             chain_type_kwargs={'prompt': PROMPT})

    if chain:
        print("Chain created")
    return chain

def answer_question(chain, question):
    time_start = time.time()
    output = chain({'query': question})
    response = output["result"]
    time_elapsed = time.time() - time_start
    print(f'response time: {time_elapsed:.02f} sec')
  
    if "Answer:" in response:
        response = response.split("Answer:")[1].strip()

    response_lines = response.split('\n')
    cleaned_response = " ".join(line.strip() for line in response_lines if not line.startswith("Question:") and not line.startswith("Document:"))

    answer = cleaned_response.strip()
    return response

if __name__ == "__main__":
    data_path = "/kaggle/input/axisfaq-data"
    documents = load_all_files(data_path)
    texts = interpret_files(documents)
    embeddings = create_embeddings()
    save(texts, embeddings)

    # model_path = "D:\Axis-FAQ-chatbot\models\llama-2-7b-chat.ggmlv3.q8_0.bin"
    model_name = input("Enter the model name (phi3 or llama or gemma): ")
    llm = load_llm(model_name)
    QA_LLM = retrieve_docs(embeddings, llm)

    while True:
        user_input = input("\n What is your question? \n")
        if user_input.lower() == "quit":
            break
        answer = answer_question(QA_LLM, user_input)
        print(f"Answer: {answer}")

        continue_input = input("\n Do you want to ask another question? (yes to continue, quit to exit, change to switch model): ").strip().lower()
        if continue_input == "quit":
            break
        elif continue_input == "change":
            model_name = input("\n Enter the new model name (phi3 or llama or gemma): ")
            llm = load_llm(model_name)
            QA_LLM = retrieve_docs(embeddings, llm)