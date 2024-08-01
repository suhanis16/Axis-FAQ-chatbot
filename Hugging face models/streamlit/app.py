import streamlit as st
import warnings
warnings.filterwarnings("ignore")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from langchain import HuggingFacePipeline
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
import torch
import pandas as pd
import pathlib
import docx
from langchain.docstore.document import Document
import os
import time
from accelerate import init_empty_weights, load_checkpoint_and_dispatch
import bitsandbytes as bnb

def clear_memory():
    torch.cuda.empty_cache()

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def load_all_files(directory_path):
    data = []
    for file_path in pathlib.Path(directory_path).glob("*"):
        if file_path.suffix == '.csv':
            df = pd.read_csv(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
        elif file_path.suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                data.append(Document(page_content=content))
        elif file_path.suffix == '.docx':
            content = read_docx(file_path)
            data.append(Document(page_content=content))
        elif file_path.suffix == '.xlsx':
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
    return data

def interpret_files(documents):
    print(f"Total documents loaded: {len(documents)}")
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    texts = splitter.split_documents(documents)
    return texts

def create_embeddings():
    print("Creating embeddings")
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cuda'}
    )
    return embeddings

def save(texts, embeddings):
    print("Saving data")
    db=Chroma.from_documents(texts,embedding=embeddings,persist_directory="test_index")
    db.persist()

def load_llm(model_name):
    model_name = model_name.lower()
    print("Loading LLM")
    if model_name == "phi3" or model_name == "phi-3":
        print("Loading Phi-3 model")
        tokenizer = AutoTokenizer.from_pretrained("microsoft/Phi-3-mini-4k-instruct")
        model = AutoModelForCausalLM.from_pretrained("microsoft/Phi-3-mini-4k-instruct", device_map='auto', torch_dtype="auto", trust_remote_code=True,)

    elif model_name == "llama":
        print("Loading Llama model")
        tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-2-7b-chat-hf")
        model = AutoModelForCausalLM.from_pretrained("meta-llama/Llama-2-7b-chat-hf", device_map='auto', torch_dtype=torch.float16, trust_remote_code=True,)

    elif model_name == "gemma":
        print("Loading GEMMA model")
        model_name = "google/gemma-2-9b-it"
        print("Loading Gemma model")

        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModelForCausalLM.from_pretrained(
            model_name,
            device_map='auto',
            load_in_8bit=True,
            torch_dtype=torch.float16,
            trust_remote_code=True
        )

    pipe = pipeline("text-generation", model=model, tokenizer=tokenizer, max_new_tokens=300)
    llm = HuggingFacePipeline(pipeline=pipe)
    return llm

def retrieve_docs(embeddings, llm):
    print("Retrieving documents")

    vectordb = Chroma(persist_directory="test_index", embedding_function=embeddings)
    retriever = vectordb.as_retriever(search_kwargs={"k": 2})

    print("Retrieved documents")

    qna_prompt_template = """Use the following pieces of information to answer the user's question. 
    If the provided context does not contain the answer, use your general knowledge to provide a helpful response.
    Context: {context}
    Question: {question}
    Answer:"""

    PROMPT = PromptTemplate(
       template=qna_prompt_template, input_variables=["context","question"]
    )

    print("Sending the chain")
    chain = RetrievalQA.from_chain_type(llm=llm,
                                        chain_type='stuff',
                                        retriever=retriever,
                                        return_source_documents=True,
                                        chain_type_kwargs={'prompt': PROMPT})

    if chain:
        print("Chain created")
    return chain

def answer_question(chain, question, memory):
    time_start = time.time()
    output = chain({'query': question})
    response = output["result"]
    time_elapsed = time.time() - time_start
    print(f'response time: {time_elapsed:.02f} sec')

    if "Answer:" in response:
        response = response.split("Answer:")[1].strip()

    response_lines = response.split('\n')
    cleaned_response = " ".join(line.strip() for line in response_lines if not line.startswith("Question:") and not line.startswith("Document:"))

    answer = cleaned_response.strip()

    memory.append((question, answer))
    return answer

def main():
    st.markdown(
        """
        <style>
        .main {
            background-color: #AE275F;
        }
        .sidebar .sidebar-content {
            background-color: #AE275F;
        }
        .header {
            text-align: center;
            padding: 10px;
            background-color: #AE275F;
            color: white;
        }
        .header img {
            max-width: 100px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.image("/content/Axis_logo.jpg", width=100)
    st.markdown(
        """
        <div class="header">
            <h1>Axis Bank FAQ Chatbot</h1>
        </div>
        """,
        unsafe_allow_html=True
    )

    st.sidebar.title("MODEL MENU")

    if 'model_name' not in st.session_state:
        st.session_state.model_name = None

    model_options = ["None", "Phi-3", "Llama", "Gemma"]

    if st.session_state.model_name is None:
        model_name = st.sidebar.selectbox("Choose a model:", model_options, key='model_name_select')
        if model_name != "None":
            st.session_state.model_name = model_name
    else:
        st.sidebar.selectbox("Choose a model:", model_options, key='model_name_select', index=model_options.index(st.session_state.model_name), disabled=True)

    st.title("Chatbot")

    if 'questions' not in st.session_state:
        st.session_state.questions = []
        st.session_state.answers = []
        st.session_state.llm = None
        st.session_state.chain = None
        st.session_state.memory = []
        st.session_state.quit = False
        st.session_state.loading = False
        st.session_state.processing = False

    if st.session_state.quit:
        st.write("You have exited the conversation.")
        return

    if st.session_state.llm is None and st.session_state.model_name:
        with st.spinner('Loading model...'):
            st.session_state.loading = True
            data_path = "/content/"
            documents = load_all_files(data_path)
            texts = interpret_files(documents)
            embeddings = create_embeddings()
            save(texts, embeddings)

            st.session_state.llm = load_llm(st.session_state.model_name)
            st.session_state.chain = retrieve_docs(embeddings, st.session_state.llm)
            st.session_state.loading = False
            st.success("Model loaded and ready for questions!")

    disable_input = st.session_state.llm is None or st.session_state.loading or st.session_state.processing

    for i in range(len(st.session_state.questions)):
        st.text_area(f"Question {i + 1}", st.session_state.questions[i], key=f"question_{i}", disabled=True)
        st.text_area(f"Answer {i + 1}", st.session_state.answers[i], key=f"answer_{i}", disabled=True)

    question = st.text_input("Ask a new question:", key="new_question", disabled=disable_input)
    if st.button("Submit", disabled=disable_input):
        if question and st.session_state.chain:
            st.session_state.processing = True
            with st.spinner('Finding the answer...'):
                answer = answer_question(st.session_state.chain, question, st.session_state.memory)
                st.session_state.questions.append(question)
                st.session_state.answers.append(answer)
                st.session_state.processing = False
                st.experimental_rerun()

    if st.button("Clear Memory", disabled=st.session_state.processing):
        st.session_state.memory.clear()
        st.success("Memory cleared.")
        st.experimental_rerun()

    if st.button("Quit", disabled=st.session_state.processing):
        st.session_state.quit = True
        st.experimental_rerun()

if __name__ == "__main__":
    main()
