import warnings
warnings.filterwarnings("ignore")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import CTransformers
from langchain_core.prompts import PromptTemplate
from langchain.chains import RetrievalQA
import time
import pathlib
import pandas as pd
import docx
from langchain.docstore.document import Document

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def load_all_files(directory_path):
    data = []
    # print(f"Loading files from directory: {directory_path}")
    for file_path in pathlib.Path(directory_path).glob("*"):
        # print(f"Found file: {file_path}")
        if file_path.suffix == '.csv':
            df = pd.read_csv(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
        elif file_path.suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                data.append(Document(page_content=content))
        elif file_path.suffix == '.docx':
            content = read_docx(file_path)
            data.append(Document(page_content=content))
        elif file_path.suffix == '.xlsx':
            df = pd.read_excel(file_path)
            for _, row in df.iterrows():
                content = " ".join(str(value) for value in row.values)
                data.append(Document(page_content=content))
    # print(f"Total documents loaded: {len(data)}")
    return data

def interpret_files(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=50)
    texts = splitter.split_documents(documents)
    # print(f"Total texts generated: {len(texts)}")
    return texts

def create_embeddings():
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )
    return embeddings

def save(texts, embeddings):
    db = FAISS.from_documents(texts, embeddings)
    db.save_local("faiss")

def load_llm(model_name, model_path):
    if model_name == "llama":
        llm = CTransformers(model=model_path,
                            model_type='llama',
                            config={'max_new_tokens': 256, 'temperature': 0.01})
    return llm

def retrieve_docs(embeddings, llm):
    template = """Use the following pieces of information to answer the user's question.
    If you don't know the answer, just say that you don't know, don't try to make up an answer.
    Context: {context}
    Question: {question}
    Only return the helpful answer below and nothing else.
    Helpful answer:
    """

    db = FAISS.load_local("faiss", embeddings, allow_dangerous_deserialization=True)
    retriever = db.as_retriever(search_kwargs={'k': 2})
    prompt = PromptTemplate(
        template=template,
        input_variables=['context', 'question'])

    QA_LLM = RetrievalQA.from_chain_type(llm=llm,
                                         chain_type='stuff',
                                         retriever=retriever,
                                         return_source_documents=True,
                                         chain_type_kwargs={'prompt': prompt})
    return QA_LLM

def query(model, question):
    time_start = time.time()
    output = model({'query': question})
    response = output["result"]
    time_elapsed = time.time() - time_start
    print(f'response time: {time_elapsed:.02f} sec')
    print(f'Question: {question}')
    print(f'Answer: \n {response}')

if __name__ == "__main__":
    data_path = "D:\Axis-FAQ-chatbot\Data"
    documents = load_all_files(data_path)
    texts = interpret_files(documents)
    embeddings = create_embeddings()

    save(texts, embeddings)

    model_path = "D:\Axis-FAQ-chatbot\models\llama-2-7b-chat.ggmlv3.q8_0.bin"

    llm = load_llm("llama", model_path)
    QA_LLM = retrieve_docs(embeddings, llm)

    user_input = input("What is your question? \n")
    query(QA_LLM, user_input)